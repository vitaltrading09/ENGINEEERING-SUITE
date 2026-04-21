"""
screen_report_widget.py
-----------------------
Screen Capture & Report Builder.

Features
--------
- Select any connected monitor from a drop-down.
- Capture Full Screen  — instant, no delay.
- Select Region        — drag a rectangle over exactly what you need.
- Minimize & Capture   — hides the app 1.5 s then grabs the full screen
                         so the Engineering Suite is never in the shot.
- Images are saved as lossless PNG at the screen's native physical resolution
  immediately to the chosen backup folder.
- Export a professional PDF or Word report: cover page, table of contents,
  date, and a three-row sign-off block, with one captioned image per page.

Layout (3 tabs)
---------------
Tab 1 — Setup          : report metadata + save folder (fill once)
Tab 2 — Capture        : prominent "NEXT" banner + screen selector + capture buttons
Tab 3 — Gallery & Export: placeholder/captured card list + export controls
"""

import datetime
import json
from pathlib import Path

from PyQt6.QtWidgets import (
    QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,
    QComboBox, QLineEdit, QScrollArea, QFileDialog, QFrame,
    QApplication, QGroupBox, QMessageBox, QDialog, QDialogButtonBox,
    QSpinBox, QPlainTextEdit, QListWidget, QTabWidget, QSizePolicy,
)
from PyQt6.QtCore import Qt, QTimer, QRect, pyqtSignal
from PyQt6.QtGui import (
    QPixmap, QFont, QPainter, QPen, QColor,
)
import qtawesome as qta

from calculators.base_calculator import BaseCalculator


# ─────────────────────────────────────────────────────────────────────────────
# Template system — built-in templates, storage, and caption generation
# ─────────────────────────────────────────────────────────────────────────────

_TEMPLATES_FILE = Path.home() / ".engineering_suite" / "screen_templates.json"

_BUILTIN_TEMPLATES: dict = {
    "Inverters (Sungrow)": {
        "title":       "Inverters",
        "unit_prefix": "INV",
        "unit_pad":    2,
        "preamble":    ["General Info"],
        "per_unit": [
            "Realtime values",
            "DC Info",
            "DC Info 02",
            "Initial Parameters",
            "Operation Parameters",
            "Operation Parameters 02",
            "Operation Parameters 03",
            "System Parameters",
            "Protection Parameters",
            "Protection Parameters 02",
            "Protection Parameters Others",
            "Protection Parameters Others 02",
            "Power Regulation",
        ],
    },
}


def _load_user_templates() -> dict:
    try:
        if _TEMPLATES_FILE.exists():
            return json.loads(_TEMPLATES_FILE.read_text("utf-8"))
    except Exception:
        pass
    return {}


def _save_user_templates(templates: dict) -> None:
    try:
        _TEMPLATES_FILE.parent.mkdir(parents=True, exist_ok=True)
        _TEMPLATES_FILE.write_text(json.dumps(templates, indent=2), "utf-8")
    except Exception:
        pass


def _all_templates() -> dict:
    merged = dict(_BUILTIN_TEMPLATES)
    merged.update(_load_user_templates())
    return merged


def _generate_captions(tpl: dict, unit_count: int) -> list:
    """Build the full ordered caption list from a template definition."""
    captions = list(tpl.get("preamble", []))
    prefix   = tpl.get("unit_prefix", "UNIT")
    pad      = int(tpl.get("unit_pad", 2))
    for i in range(1, unit_count + 1):
        uid = f"{prefix} {str(i).zfill(pad)}"
        for cap in tpl.get("per_unit", []):
            captions.append(f"{uid}:{cap}")
    return captions


# ─────────────────────────────────────────────────────────────────────────────
# Region selection overlay
# ─────────────────────────────────────────────────────────────────────────────

class _RegionOverlay(QWidget):
    """
    Fullscreen overlay shown on the target monitor.

    The screen is grabbed *before* this window opens so the overlay itself
    never appears in the final capture.  The screenshot is rendered as a
    background so the user can see exactly what they are selecting.

    Emits
    -----
    region_selected(QRect)  — logical-pixel rectangle on success
    cancelled()             — on Escape or a click too small to be useful
    """

    region_selected = pyqtSignal(object)   # QRect in widget-local logical px
    cancelled       = pyqtSignal()

    def __init__(self, screenshot: QPixmap, screen_geometry: QRect,
                 parent=None):
        super().__init__(parent,
                         Qt.WindowType.FramelessWindowHint |
                         Qt.WindowType.WindowStaysOnTopHint |
                         Qt.WindowType.Tool)
        self._bg     = screenshot
        self._start  = None
        self._end    = None
        self._active = False

        self.setGeometry(screen_geometry)
        self.setCursor(Qt.CursorShape.CrossCursor)
        self.setMouseTracking(True)

    # ── Events ────────────────────────────────────────────────────────────────

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Escape:
            self.cancelled.emit()
            self.close()

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._start  = event.pos()
            self._end    = event.pos()
            self._active = True
            self.update()

    def mouseMoveEvent(self, event):
        if self._active:
            self._end = event.pos()
            self.update()

    def mouseReleaseEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton and self._active:
            self._active = False
            self._end    = event.pos()
            if self._start:
                rect = QRect(self._start, self._end).normalized()
                if rect.width() > 4 and rect.height() > 4:
                    self.region_selected.emit(rect)
                    self.close()
                    return
            self.cancelled.emit()
            self.close()

    # ── Paint ─────────────────────────────────────────────────────────────────

    def paintEvent(self, _event):
        p = QPainter(self)

        # 1. Draw the pre-captured screenshot as background (1:1 logical mapping)
        p.drawPixmap(self.rect(), self._bg,
                     QRect(0, 0, self._bg.width(), self._bg.height()))

        # 2. Semi-transparent dark veil over everything
        p.fillRect(self.rect(), QColor(0, 0, 0, 115))

        if self._start and self._end:
            sel = QRect(self._start, self._end).normalized()

            # 3. Cut out the selection — show original screenshot through it
            p.drawPixmap(sel, self._bg, sel)

            # 4. Bright green border around selection
            pen = QPen(QColor("#2ea043"), 2)
            pen.setStyle(Qt.PenStyle.SolidLine)
            p.setPen(pen)
            p.setBrush(Qt.BrushStyle.NoBrush)
            p.drawRect(sel)

            # 5. Dimension label
            label = f"  {sel.width()} × {sel.height()} px  "
            font  = QFont("Segoe UI", 10, QFont.Weight.Bold)
            p.setFont(font)
            fm  = p.fontMetrics()
            lw  = fm.horizontalAdvance(label)
            lh  = fm.height()
            lx  = sel.x()
            ly  = sel.y() - lh - 4 if sel.y() > lh + 4 else sel.bottom() + lh + 4
            p.fillRect(lx, ly - lh, lw, lh + 4, QColor(0, 0, 0, 180))
            p.setPen(QColor("white"))
            p.drawText(lx, ly, label)

        p.end()


# ─────────────────────────────────────────────────────────────────────────────
# Template dialogs
# ─────────────────────────────────────────────────────────────────────────────

class _TemplateEditDialog(QDialog):
    """Create or edit a caption template (name, prefix, preamble, per-unit list)."""

    def __init__(self, name: str = "", data: dict | None = None, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Edit Template" if name else "New Template")
        self.setMinimumSize(500, 580)
        self._build(name, data or {})

    def _build(self, name: str, data: dict):
        lay = QVBoxLayout(self)
        lay.setSpacing(10)

        def lf(text):
            l = QLabel(text)
            l.setObjectName("header_subtitle")
            return l

        # Template name
        r1 = QHBoxLayout()
        r1.addWidget(lf("Template Name:"))
        self._name_edit = QLineEdit(name)
        r1.addWidget(self._name_edit, 1)
        lay.addLayout(r1)

        # Default report title
        r2 = QHBoxLayout()
        r2.addWidget(lf("Default Report Title:"))
        self._title_edit = QLineEdit(data.get("title", ""))
        r2.addWidget(self._title_edit, 1)
        lay.addLayout(r2)

        # Unit prefix + zero-pad
        r3 = QHBoxLayout()
        r3.addWidget(lf("Unit Prefix  (e.g. INV, MDB, TRF):"))
        self._prefix_edit = QLineEdit(data.get("unit_prefix", "INV"))
        self._prefix_edit.setFixedWidth(90)
        r3.addWidget(self._prefix_edit)
        r3.addWidget(lf("   Zero-pad digits:"))
        self._pad_spin = QSpinBox()
        self._pad_spin.setRange(1, 4)
        self._pad_spin.setValue(int(data.get("unit_pad", 2)))
        self._pad_spin.setFixedWidth(60)
        r3.addWidget(self._pad_spin)
        r3.addStretch()
        lay.addLayout(r3)

        # Preamble
        lay.addWidget(lf(
            "Preamble captions — appear once before the repeating block (one per line):"))
        self._preamble_edit = QPlainTextEdit(
            "\n".join(data.get("preamble", [])))
        self._preamble_edit.setMaximumHeight(70)
        self._preamble_edit.setPlaceholderText("e.g. General Info")
        lay.addWidget(self._preamble_edit)

        # Per-unit captions
        lay.addWidget(lf(
            "Per-unit captions — repeated for every unit (one per line):"))
        self._perunit_edit = QPlainTextEdit(
            "\n".join(data.get("per_unit", [])))
        self._perunit_edit.setPlaceholderText(
            "e.g.\nRealtime values\nDC Info\nDC Info 02\nInitial Parameters")
        lay.addWidget(self._perunit_edit, 1)

        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Cancel |
            QDialogButtonBox.StandardButton.Ok)
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("Save Template")
        btns.accepted.connect(self._on_accept)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)

    def _on_accept(self):
        if not self._name_edit.text().strip():
            QMessageBox.warning(self, "Validation",
                                "Template name cannot be empty.")
            return
        self.accept()

    def get_template(self) -> tuple:
        """Return ``(name, data_dict)``."""
        name     = self._name_edit.text().strip()
        preamble = [l.strip() for l in
                    self._preamble_edit.toPlainText().splitlines() if l.strip()]
        per_unit = [l.strip() for l in
                    self._perunit_edit.toPlainText().splitlines()  if l.strip()]
        return name, {
            "title":       self._title_edit.text().strip(),
            "unit_prefix": self._prefix_edit.text().strip() or "UNIT",
            "unit_pad":    self._pad_spin.value(),
            "preamble":    preamble,
            "per_unit":    per_unit,
        }


class _TemplateDialog(QDialog):
    """
    Choose a caption template, set the number of units, preview the
    full generated caption list, and manage (create / edit / delete)
    custom templates.
    """

    def __init__(self, current_queue_len: int = 0, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Load Caption Template")
        self.setMinimumSize(580, 540)
        self._captions: list = []
        self._build(current_queue_len)

    def _build(self, current_queue_len: int):
        lay = QVBoxLayout(self)
        lay.setSpacing(8)

        def lf(text):
            l = QLabel(text)
            l.setObjectName("header_subtitle")
            return l

        # ── Template selector ─────────────────────────────────────────────────
        r1 = QHBoxLayout()
        r1.addWidget(lf("Template:"))
        self._tpl_combo = QComboBox()
        self._tpl_combo.setMinimumWidth(220)
        r1.addWidget(self._tpl_combo, 1)
        new_btn  = QPushButton("+ New")
        edit_btn = QPushButton("Edit")
        del_btn  = QPushButton("Delete")
        for b in (new_btn, edit_btn, del_btn):
            b.setObjectName("secondary_btn")
            b.setFixedHeight(28)
            r1.addWidget(b)
        new_btn.clicked.connect(self._new_template)
        edit_btn.clicked.connect(self._edit_template)
        del_btn.clicked.connect(self._delete_template)
        lay.addLayout(r1)

        # ── Unit count ────────────────────────────────────────────────────────
        r2 = QHBoxLayout()
        self._unit_lbl = lf("Number of units:")
        r2.addWidget(self._unit_lbl)
        self._unit_spin = QSpinBox()
        self._unit_spin.setRange(1, 200)
        self._unit_spin.setValue(13)
        self._unit_spin.setFixedWidth(90)
        r2.addWidget(self._unit_spin)
        r2.addStretch()
        lay.addLayout(r2)

        # ── Preview list ──────────────────────────────────────────────────────
        self._preview_lbl = lf("Preview (0 captions):")
        lay.addWidget(self._preview_lbl)
        self._preview = QListWidget()
        self._preview.setMinimumHeight(220)
        self._preview.setStyleSheet(
            "QListWidget { background: #0d1117; color: #c9d1d9;"
            " border: 1px solid #30363d; }")
        lay.addWidget(self._preview, 1)

        # ── Warning if replacing existing queue ───────────────────────────────
        if current_queue_len > 0:
            warn = QLabel(
                f"⚠  Loading a new template will replace the current "
                f"{current_queue_len}-caption queue.")
            warn.setStyleSheet("color: #f0883e; padding: 2px 0;")
            warn.setWordWrap(True)
            lay.addWidget(warn)

        # ── Dialog buttons ────────────────────────────────────────────────────
        btns = QDialogButtonBox(
            QDialogButtonBox.StandardButton.Cancel |
            QDialogButtonBox.StandardButton.Ok)
        btns.button(QDialogButtonBox.StandardButton.Ok).setText("Load Template")
        btns.accepted.connect(self.accept)
        btns.rejected.connect(self.reject)
        lay.addWidget(btns)

        # ── Populate ──────────────────────────────────────────────────────────
        self._refresh_combo()
        self._tpl_combo.currentIndexChanged.connect(self._refresh_preview)
        self._unit_spin.valueChanged.connect(self._refresh_preview)

    # ── Internal helpers ──────────────────────────────────────────────────────

    def _refresh_combo(self):
        cur = self._tpl_combo.currentText()
        self._tpl_combo.blockSignals(True)
        self._tpl_combo.clear()
        for name in _all_templates():
            self._tpl_combo.addItem(name)
        idx = self._tpl_combo.findText(cur)
        self._tpl_combo.setCurrentIndex(max(0, idx))
        self._tpl_combo.blockSignals(False)
        self._refresh_preview()

    def _current_tpl(self) -> dict | None:
        return _all_templates().get(self._tpl_combo.currentText())

    def _refresh_preview(self):
        tpl = self._current_tpl()
        self._preview.clear()
        if not tpl:
            return
        prefix = tpl.get("unit_prefix", "UNIT")
        self._unit_lbl.setText(f"Number of {prefix} units:")
        caps = _generate_captions(tpl, self._unit_spin.value())
        self._captions = caps
        self._preview_lbl.setText(f"Preview ({len(caps)} captions):")
        for i, cap in enumerate(caps, 1):
            self._preview.addItem(f"{i}. {cap}")

    def _new_template(self):
        dlg = _TemplateEditDialog(parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            name, data = dlg.get_template()
            user = _load_user_templates()
            user[name] = data
            _save_user_templates(user)
            self._refresh_combo()
            idx = self._tpl_combo.findText(name)
            if idx >= 0:
                self._tpl_combo.setCurrentIndex(idx)

    def _edit_template(self):
        name = self._tpl_combo.currentText()
        if name in _BUILTIN_TEMPLATES:
            QMessageBox.information(
                self, "Built-in Template",
                "Built-in templates cannot be edited.\n"
                "Use  '+ New'  to create a custom copy.")
            return
        tpl = _load_user_templates().get(name)
        if not tpl:
            return
        dlg = _TemplateEditDialog(name=name, data=tpl, parent=self)
        if dlg.exec() == QDialog.DialogCode.Accepted:
            new_name, new_data = dlg.get_template()
            user = _load_user_templates()
            if new_name != name and name in user:
                del user[name]
            user[new_name] = new_data
            _save_user_templates(user)
            self._refresh_combo()

    def _delete_template(self):
        name = self._tpl_combo.currentText()
        if name in _BUILTIN_TEMPLATES:
            QMessageBox.information(
                self, "Built-in Template",
                "Built-in templates cannot be deleted.")
            return
        user = _load_user_templates()
        if name not in user:
            return
        if QMessageBox.question(
            self, "Delete Template",
            f"Permanently delete  '{name}' ?",
            QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        ) == QMessageBox.StandardButton.Yes:
            del user[name]
            _save_user_templates(user)
            self._refresh_combo()

    # ── Public API ────────────────────────────────────────────────────────────

    def get_captions(self) -> list:
        return list(self._captions)


# ─────────────────────────────────────────────────────────────────────────────
# Screenshot card  (thumbnail + caption + reorder / delete)
# ─────────────────────────────────────────────────────────────────────────────

class ScreenshotCard(QFrame):
    """
    One entry in the gallery — either a captured screenshot or a pending
    placeholder created from a template.

    Placeholder cards are visually distinct (dashed border, grey thumbnail,
    "Capture pending" label).  Call ``set_captured()`` to promote a placeholder
    to a real card once the user has taken the screenshot.
    """

    sig_delete       = pyqtSignal(object)
    sig_move_up      = pyqtSignal(object)
    sig_move_down    = pyqtSignal(object)
    sig_retake       = pyqtSignal(object)   # clear capture, become placeholder again
    sig_insert_below = pyqtSignal(object)   # add a new blank slot below this card

    def __init__(self, pixmap: QPixmap | None, filepath: str, index: int,
                 caption: str = "", parent=None):
        super().__init__(parent)
        self._pixmap         = pixmap
        self.filepath        = filepath
        self._index          = index
        self._is_placeholder = (pixmap is None)
        self._build()
        if caption:
            self.caption_edit.setText(caption)
        self._apply_style()

    # ── public interface ──────────────────────────────────────────────────────

    @property
    def is_placeholder(self) -> bool:
        return self._is_placeholder

    def set_captured(self, pixmap: QPixmap, filepath: str):
        """Promote this placeholder to a captured card."""
        self._pixmap         = pixmap
        self.filepath        = filepath
        self._is_placeholder = False
        self._apply_style()
        self._refresh()

    def set_index(self, idx: int):
        self._index = idx
        self._num_lbl.setText(f"#{idx}")
        self.caption_edit.setPlaceholderText(
            f"Image {idx} — enter description…")

    def get_caption(self) -> str:
        t = self.caption_edit.text().strip()
        return t if t else f"Image {self._index}"

    # ── internal build ────────────────────────────────────────────────────────

    def _build(self):
        self.setObjectName("card")
        self.setFrameShape(QFrame.Shape.StyledPanel)

        root = QHBoxLayout(self)
        root.setContentsMargins(10, 8, 10, 8)
        root.setSpacing(14)

        # ① Index badge
        self._num_lbl = QLabel()
        self._num_lbl.setFixedWidth(30)
        self._num_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._num_lbl.setFont(QFont("Segoe UI", 11, QFont.Weight.Bold))
        root.addWidget(self._num_lbl)

        # ② Thumbnail
        self._thumb = QLabel()
        self._thumb.setFixedSize(240, 135)
        self._thumb.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._thumb.setStyleSheet(
            "border: 1px solid #30363d; border-radius: 4px; background: #0d1117;")
        root.addWidget(self._thumb)

        # ③ Right column
        right = QVBoxLayout()
        right.setSpacing(6)

        cap_lbl = QLabel("Caption / Sub-heading")
        cap_lbl.setObjectName("header_subtitle")
        right.addWidget(cap_lbl)

        self.caption_edit = QLineEdit()
        self.caption_edit.setMinimumWidth(320)
        right.addWidget(self.caption_edit)

        self._path_lbl = QLabel()
        self._path_lbl.setObjectName("header_subtitle")
        self._path_lbl.setTextInteractionFlags(
            Qt.TextInteractionFlag.TextSelectableByMouse)
        right.addWidget(self._path_lbl)

        right.addStretch()

        ctrl = QHBoxLayout()
        ctrl.setSpacing(4)
        for icon_name, tip, sig in [
            ("fa5s.arrow-up",   "Move up",   self.sig_move_up),
            ("fa5s.arrow-down", "Move down", self.sig_move_down),
        ]:
            btn = QPushButton()
            btn.setIcon(qta.icon(icon_name, color="#8b949e"))
            btn.setFixedSize(28, 28)
            btn.setToolTip(tip)
            btn.setObjectName("secondary_btn")
            btn.clicked.connect(lambda _=False, s=sig: s.emit(self))
            ctrl.addWidget(btn)
        ctrl.addStretch()

        # ↺ Retake — only visible on captured (non-placeholder) cards
        self._retake_btn = QPushButton()
        self._retake_btn.setIcon(qta.icon("fa5s.redo", color="#e3b341"))
        self._retake_btn.setFixedSize(28, 28)
        self._retake_btn.setToolTip(
            "Clear this screenshot and mark as pending so you can retake it.\n"
            "(The saved PNG on disk is NOT deleted.)")
        self._retake_btn.setObjectName("secondary_btn")
        self._retake_btn.clicked.connect(lambda: self.sig_retake.emit(self))
        ctrl.addWidget(self._retake_btn)

        # + Insert a new blank capture slot below this card
        add_btn = QPushButton()
        add_btn.setIcon(qta.icon("fa5s.plus", color="#3fb950"))
        add_btn.setFixedSize(28, 28)
        add_btn.setToolTip("Insert a new empty capture slot below this one")
        add_btn.setObjectName("secondary_btn")
        add_btn.clicked.connect(lambda: self.sig_insert_below.emit(self))
        ctrl.addWidget(add_btn)

        # 🗑 Remove card from list entirely
        del_btn = QPushButton()
        del_btn.setIcon(qta.icon("fa5s.trash", color="#f85149"))
        del_btn.setFixedSize(28, 28)
        del_btn.setToolTip("Remove this card from the list\n"
                           "(saved PNG on disk is NOT deleted)")
        del_btn.setObjectName("secondary_btn")
        del_btn.clicked.connect(lambda: self.sig_delete.emit(self))
        ctrl.addWidget(del_btn)

        right.addLayout(ctrl)

        root.addLayout(right, 1)
        self._refresh()

    def _refresh(self):
        self._num_lbl.setText(f"#{self._index}")
        if self._is_placeholder:
            # Draw a grey "Capture pending" thumbnail
            px = QPixmap(240, 135)
            px.fill(QColor("#12171e"))
            p = QPainter(px)
            p.setPen(QColor("#484f58"))
            p.setFont(QFont("Segoe UI", 9))
            p.drawText(px.rect(), Qt.AlignmentFlag.AlignCenter,
                       "Capture pending…")
            p.end()
            self._thumb.setPixmap(px)
            self._path_lbl.setText("screenshot not yet captured")
        else:
            self._thumb.setPixmap(
                self._pixmap.scaled(240, 135,
                                    Qt.AspectRatioMode.KeepAspectRatio,
                                    Qt.TransformationMode.SmoothTransformation))
            self._path_lbl.setText(f"Saved  →  {Path(self.filepath).name}")
        self.caption_edit.setPlaceholderText(
            f"Image {self._index} — enter description…")

    def reset_to_placeholder(self):
        """
        Demote this captured card back to a pending placeholder.
        The PNG that was saved to disk is NOT deleted — it just becomes
        unlinked from this slot so the user can retake the screenshot.
        """
        self._pixmap         = None
        self.filepath        = ""
        self._is_placeholder = True
        self._apply_style()
        self._refresh()

    def _apply_style(self):
        is_ph = self._is_placeholder
        if is_ph:
            self.setStyleSheet(
                "ScreenshotCard { border: 1px dashed #30363d;"
                " background: #0d1117; }")
            self._num_lbl.setStyleSheet("color: #484f58;")
        else:
            self.setStyleSheet("")
            self._num_lbl.setStyleSheet("")
        # Retake button only makes sense when there IS a capture to undo
        if hasattr(self, "_retake_btn"):
            self._retake_btn.setVisible(not is_ph)


# ─────────────────────────────────────────────────────────────────────────────
# Main widget
# ─────────────────────────────────────────────────────────────────────────────

class ScreenReportWidget(BaseCalculator):
    """
    Screen Capture & Report Builder.

    Opens maximised so the preview panel and gallery have plenty of space,
    especially when working across two monitors.
    """

    calculator_name  = "Screen Report Builder"
    sans_reference   = "Internal — Solareff Engineering"
    prefer_maximized = True   # main_window reads this to showMaximized()

    def __init__(self, parent=None):
        super().__init__(parent)
        self._cards:           list[ScreenshotCard] = []
        self._caption_queue:   list[str]            = []
        self._folder:          Path = Path.home() / "Documents"
        self._capture_counter: int  = 0
        self._session_path:    Path | None          = None
        self._live_timer             = QTimer(self)
        self._live_timer.setInterval(1000)
        self._overlay = None          # holds _RegionOverlay while active

        self._build_ui()
        self._wire_signals()
        self._refresh_screens()

    # ═══════════════════════════════════════════════════════════════════════════
    # UI — 3-tab layout
    # ═══════════════════════════════════════════════════════════════════════════

    def _build_ui(self):
        root = QVBoxLayout(self)
        root.setContentsMargins(0, 0, 0, 0)
        root.setSpacing(0)

        self._tabs = QTabWidget()
        self._tabs.setObjectName("report_tabs")

        # ── Tab 1: Setup ──────────────────────────────────────────────────────
        t1 = QWidget()
        t1_lay = QVBoxLayout(t1)
        t1_lay.setContentsMargins(20, 16, 20, 16)
        t1_lay.setSpacing(12)
        t1_lay.addWidget(self._build_settings_group())
        t1_lay.addStretch()
        self._tabs.addTab(
            t1,
            qta.icon("fa5s.cog", color="#8b949e"),
            "  Setup")

        # ── Tab 2: Capture ────────────────────────────────────────────────────
        t2 = QWidget()
        t2_lay = QVBoxLayout(t2)
        t2_lay.setContentsMargins(16, 12, 16, 12)
        t2_lay.setSpacing(10)
        t2_lay.addWidget(self._build_next_capture_banner())
        t2_lay.addWidget(self._build_capture_group(), 1)   # fills remaining tab height
        self._tabs.addTab(
            t2,
            qta.icon("fa5s.camera", color="#8b949e"),
            "  Capture")

        # ── Tab 3: Gallery & Export ───────────────────────────────────────────
        t3 = QWidget()
        t3_lay = QVBoxLayout(t3)
        t3_lay.setContentsMargins(8, 8, 8, 8)
        t3_lay.setSpacing(6)
        t3_lay.addWidget(self._build_gallery_group(), 1)
        t3_lay.addWidget(self._build_export_bar())
        self._tabs.addTab(
            t3,
            qta.icon("fa5s.images", color="#8b949e"),
            "  Gallery & Export")

        root.addWidget(self._tabs, 1)

    # ── Next-capture banner (Tab 2) ───────────────────────────────────────────

    def _build_next_capture_banner(self) -> QFrame:
        """
        Prominent status card at the top of the Capture tab.
        Shows the caption for the next pending screenshot and overall progress.
        """
        frame = QFrame()
        frame.setObjectName("capture_banner")
        frame.setStyleSheet(
            "QFrame#capture_banner {"
            "  background: #161b22;"
            "  border: 1px solid #21262d;"
            "  border-radius: 8px;"
            "}")

        lay = QVBoxLayout(frame)
        lay.setContentsMargins(20, 14, 20, 14)
        lay.setSpacing(5)

        self._next_lbl = QLabel("Ready to capture")
        self._next_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._next_lbl.setFont(QFont("Segoe UI", 14, QFont.Weight.Bold))
        self._next_lbl.setStyleSheet("color: #8b949e;")
        lay.addWidget(self._next_lbl)

        self._progress_lbl = QLabel("Load a template or capture freely")
        self._progress_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._progress_lbl.setObjectName("header_subtitle")
        lay.addWidget(self._progress_lbl)

        return frame

    # ── Report settings (Tab 1) ────────────────────────────────────────────────

    def _build_settings_group(self) -> QGroupBox:
        grp = QGroupBox("Report Settings")
        g   = QVBoxLayout(grp)
        g.setSpacing(8)

        r1 = QHBoxLayout()
        self.title_edit   = self._lf("Report Title", r1)
        self.date_edit    = self._lf("Date", r1,
                                     default=datetime.date.today()
                                     .strftime("%d %B %Y"))
        g.addLayout(r1)

        r2 = QHBoxLayout()
        self.project_edit = self._lf("Project No.",   r2)
        self.client_edit  = self._lf("Client / Site", r2)
        g.addLayout(r2)

        r3 = QHBoxLayout()
        self.prepby_edit   = self._lf("Prepared By", r3)
        self.location_edit = self._lf("Location",    r3)
        g.addLayout(r3)

        r4 = QHBoxLayout()
        r4.setSpacing(6)
        lbl = QLabel("Save Folder:")
        lbl.setObjectName("header_subtitle")
        r4.addWidget(lbl)
        self.folder_edit = QLineEdit(str(self._folder))
        self.folder_edit.setReadOnly(True)
        r4.addWidget(self.folder_edit, 1)
        browse = QPushButton("Browse…")
        browse.setObjectName("secondary_btn")
        browse.setFixedHeight(30)
        browse.clicked.connect(self._browse_folder)
        r4.addWidget(browse)
        g.addLayout(r4)

        return grp

    @staticmethod
    def _lf(label: str, layout: QHBoxLayout,
            default: str = "") -> QLineEdit:
        col = QVBoxLayout()
        lbl = QLabel(label)
        lbl.setObjectName("header_subtitle")
        col.addWidget(lbl)
        ed = QLineEdit(default)
        ed.setMinimumWidth(160)
        col.addWidget(ed)
        layout.addLayout(col)
        return ed

    # ── Capture panel (Tab 2) ──────────────────────────────────────────────────

    def _build_capture_group(self) -> QGroupBox:
        grp = QGroupBox("Screen Capture")
        g   = QVBoxLayout(grp)
        g.setSpacing(10)

        # Screen selector row
        sel = QHBoxLayout()
        sel.setSpacing(8)
        lbl = QLabel("Screen:")
        lbl.setObjectName("header_subtitle")
        sel.addWidget(lbl)
        self.screen_combo = QComboBox()
        self.screen_combo.setMinimumWidth(300)
        sel.addWidget(self.screen_combo)
        ref_btn = QPushButton()
        ref_btn.setIcon(qta.icon("fa5s.sync-alt", color="#8b949e"))
        ref_btn.setFixedSize(28, 28)
        ref_btn.setToolTip("Refresh screen list")
        ref_btn.setObjectName("secondary_btn")
        ref_btn.clicked.connect(self._refresh_screens)
        sel.addWidget(ref_btn)
        sel.addStretch()
        g.addLayout(sel)

        # Preview + buttons side by side
        mid = QHBoxLayout()
        mid.setSpacing(16)

        # Preview pane — expands to fill all available horizontal and vertical space
        pv = QVBoxLayout()
        pv_lbl = QLabel("Live Preview")
        pv_lbl.setObjectName("header_subtitle")
        pv.addWidget(pv_lbl)
        self.preview_lbl = QLabel("Select a screen above")
        self.preview_lbl.setMinimumSize(480, 270)
        self.preview_lbl.setSizePolicy(
            QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.preview_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.preview_lbl.setStyleSheet(
            "border: 1px solid #30363d; border-radius: 4px;"
            " background: #0d1117; color: #8b949e;")
        pv.addWidget(self.preview_lbl, 1)
        mid.addLayout(pv, 1)   # preview gets all extra horizontal space

        # Capture buttons (fixed width on the right)
        bc = QVBoxLayout()
        bc.setSpacing(10)
        bc.setAlignment(Qt.AlignmentFlag.AlignTop)

        self.capture_btn = QPushButton("  Capture Full Screen")
        self.capture_btn.setIcon(qta.icon("fa5s.camera", color="white"))
        self.capture_btn.setMinimumHeight(42)
        self.capture_btn.setObjectName("primary_btn")
        bc.addWidget(self.capture_btn)

        self.region_btn = QPushButton("  Select Region")
        self.region_btn.setIcon(qta.icon("fa5s.crop-alt", color="white"))
        self.region_btn.setMinimumHeight(42)
        self.region_btn.setObjectName("primary_btn")
        bc.addWidget(self.region_btn)

        self.minimize_btn = QPushButton("  Minimize & Capture")
        self.minimize_btn.setIcon(qta.icon("fa5s.window-minimize", color="white"))
        self.minimize_btn.setMinimumHeight(42)
        self.minimize_btn.setObjectName("primary_btn")
        bc.addWidget(self.minimize_btn)

        self.live_btn = QPushButton("  Live Preview (1 fps)")
        self.live_btn.setIcon(qta.icon("fa5s.eye", color="#8b949e"))
        self.live_btn.setMinimumHeight(32)
        self.live_btn.setObjectName("secondary_btn")
        self.live_btn.setCheckable(True)
        bc.addWidget(self.live_btn)

        hint = QLabel(
            "Tip: use  Minimize & Capture  to\n"
            "grab another window without the\n"
            "Engineering Suite in the shot.")
        hint.setObjectName("header_subtitle")
        hint.setWordWrap(True)
        bc.addWidget(hint)
        bc.addStretch()

        mid.addLayout(bc)
        g.addLayout(mid, 1)   # mid fills the group box vertically
        return grp

    # ── Gallery (Tab 3) ────────────────────────────────────────────────────────

    def _build_gallery_group(self) -> QGroupBox:
        grp = QGroupBox("Captured Screenshots")
        g   = QVBoxLayout(grp)
        g.setSpacing(4)

        hdr_row = QHBoxLayout()
        hdr_row.setSpacing(8)

        self._count_lbl = QLabel("0 screenshots")
        self._count_lbl.setObjectName("header_subtitle")
        hdr_row.addWidget(self._count_lbl)

        hdr_row.addStretch()

        tpl_btn = QPushButton("  Templates…")
        tpl_btn.setIcon(qta.icon("fa5s.list-alt", color="#8b949e"))
        tpl_btn.setFixedHeight(28)
        tpl_btn.setObjectName("secondary_btn")
        tpl_btn.clicked.connect(self._show_template_dialog)
        hdr_row.addWidget(tpl_btn)

        g.addLayout(hdr_row)

        self._scroll = QScrollArea()
        self._scroll.setWidgetResizable(True)
        self._scroll.setMinimumHeight(280)
        self._scroll.setHorizontalScrollBarPolicy(
            Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self._scroll.setStyleSheet(
            "QScrollArea { border: none; background: transparent; }")

        self._gallery_widget = QWidget()
        self._gallery_layout = QVBoxLayout(self._gallery_widget)
        self._gallery_layout.setContentsMargins(0, 0, 4, 0)
        self._gallery_layout.setSpacing(6)
        self._gallery_layout.addStretch()   # always last

        self._empty_lbl = QLabel(
            "No screenshots yet — use the  Capture  tab to begin,\n"
            "or load a  Template  to create a capture checklist.")
        self._empty_lbl.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self._empty_lbl.setObjectName("header_subtitle")
        self._gallery_layout.insertWidget(0, self._empty_lbl)

        self._scroll.setWidget(self._gallery_widget)
        g.addWidget(self._scroll, 1)
        return grp

    # ── Export bar (Tab 3 bottom) ──────────────────────────────────────────────

    def _build_export_bar(self) -> QFrame:
        bar = QFrame()
        lay = QHBoxLayout(bar)
        lay.setContentsMargins(0, 4, 0, 0)
        lay.setSpacing(8)

        self.pdf_btn = QPushButton("  Export PDF Report")
        self.pdf_btn.setIcon(qta.icon("fa5s.file-pdf", color="white"))
        self.pdf_btn.setMinimumHeight(38)
        self.pdf_btn.setObjectName("primary_btn")

        self.word_btn = QPushButton("  Export Word Report")
        self.word_btn.setIcon(qta.icon("fa5s.file-word", color="white"))
        self.word_btn.setMinimumHeight(38)
        self.word_btn.setObjectName("primary_btn")

        self.save_session_btn = QPushButton("  Save Session")
        self.save_session_btn.setIcon(qta.icon("fa5s.save", color="#8b949e"))
        self.save_session_btn.setMinimumHeight(38)
        self.save_session_btn.setObjectName("secondary_btn")
        self.save_session_btn.setToolTip(
            "Save the current session (cards, captions, settings) to a .srp file\n"
            "so you can close the app and resume later.")

        self.load_session_btn = QPushButton("  Load Session")
        self.load_session_btn.setIcon(qta.icon("fa5s.folder-open", color="#8b949e"))
        self.load_session_btn.setMinimumHeight(38)
        self.load_session_btn.setObjectName("secondary_btn")
        self.load_session_btn.setToolTip(
            "Open a previously saved .srp session file.")

        self.clear_btn = QPushButton("  Clear All")
        self.clear_btn.setIcon(qta.icon("fa5s.trash", color="#f85149"))
        self.clear_btn.setMinimumHeight(38)
        self.clear_btn.setObjectName("secondary_btn")

        lay.addWidget(self.pdf_btn)
        lay.addWidget(self.word_btn)
        lay.addStretch()
        lay.addWidget(self.save_session_btn)
        lay.addWidget(self.load_session_btn)
        lay.addWidget(self.clear_btn)
        return bar

    # ═══════════════════════════════════════════════════════════════════════════
    # SIGNALS
    # ═══════════════════════════════════════════════════════════════════════════

    def _wire_signals(self):
        self.capture_btn.clicked.connect(self._do_capture_full)
        self.region_btn.clicked.connect(self._start_region_capture)
        self.minimize_btn.clicked.connect(self._minimize_capture)
        self.live_btn.toggled.connect(self._toggle_live)
        self.screen_combo.currentIndexChanged.connect(self._snap_preview)
        self._live_timer.timeout.connect(self._snap_preview)
        self.pdf_btn.clicked.connect(self._export_pdf)
        self.word_btn.clicked.connect(self._export_word)
        self.clear_btn.clicked.connect(self._clear_all)
        self.save_session_btn.clicked.connect(self._save_session)
        self.load_session_btn.clicked.connect(self._load_session)

    # ═══════════════════════════════════════════════════════════════════════════
    # SCREEN HELPERS
    # ═══════════════════════════════════════════════════════════════════════════

    def _refresh_screens(self):
        self.screen_combo.blockSignals(True)
        self.screen_combo.clear()
        for i, s in enumerate(QApplication.screens()):
            g = s.geometry()
            self.screen_combo.addItem(
                f"Screen {i + 1}  —  {g.width()} × {g.height()}  ({s.name()})")
        self.screen_combo.blockSignals(False)
        self._snap_preview()

    def _current_screen(self):
        idx     = self.screen_combo.currentIndex()
        screens = QApplication.screens()
        return screens[idx] if 0 <= idx < len(screens) \
            else QApplication.primaryScreen()

    # ── PIL-based capture helpers ──────────────────────────────────────────────
    # Qt's grabWindow() uses the Windows compositor and returns a black image
    # for non-primary monitors.  PIL's ImageGrab calls GDI BitBlt directly
    # and works reliably on every monitor.

    @staticmethod
    def _pil_grab(x1: int, y1: int, x2: int, y2: int):
        """Capture a rectangle of the virtual desktop via PIL ImageGrab/GDI."""
        from PIL import ImageGrab
        return ImageGrab.grab(bbox=(x1, y1, x2, y2), all_screens=True)

    @staticmethod
    def _pil_to_pixmap(pil_img) -> QPixmap:
        """Convert a PIL Image to QPixmap via in-memory PNG bytes."""
        import io
        buf = io.BytesIO()
        pil_img.save(buf, format="PNG")
        px = QPixmap()
        px.loadFromData(buf.getvalue())
        return px

    def _screen_bbox(self):
        """Return the virtual-desktop bounding box of the selected screen."""
        geom = self._current_screen().geometry()
        return (geom.x(), geom.y(),
                geom.x() + geom.width(),
                geom.y() + geom.height())

    def _snap_preview(self):
        try:
            bbox = self._screen_bbox()
            pil  = self._pil_grab(*bbox)
            # Scale to the label's actual painted size (grows with the window).
            w = max(480, self.preview_lbl.width())
            h = max(270, self.preview_lbl.height())
            self.preview_lbl.setPixmap(
                self._pil_to_pixmap(pil).scaled(
                    w, h,
                    Qt.AspectRatioMode.KeepAspectRatio,
                    Qt.TransformationMode.SmoothTransformation))
        except Exception:
            pass

    # ═══════════════════════════════════════════════════════════════════════════
    # CAPTURE — FULL SCREEN  (instant, no delay)
    # ═══════════════════════════════════════════════════════════════════════════

    def _do_capture_full(self):
        self._save_card(self._pil_grab(*self._screen_bbox()))

    # ═══════════════════════════════════════════════════════════════════════════
    # CAPTURE — REGION  (overlay drag-select)
    # ═══════════════════════════════════════════════════════════════════════════

    def _start_region_capture(self):
        screen = self._current_screen()
        geom   = screen.geometry()
        bbox   = (geom.x(), geom.y(),
                  geom.x() + geom.width(),
                  geom.y() + geom.height())

        # Grab BEFORE the overlay appears — overlay must never be in the shot.
        pil_full = self._pil_grab(*bbox)
        full_px  = self._pil_to_pixmap(pil_full)

        self._overlay = _RegionOverlay(full_px, geom)
        self._overlay.region_selected.connect(
            lambda rect: self._finish_region_capture(rect, pil_full))
        self._overlay.cancelled.connect(
            lambda: setattr(self, "_overlay", None))
        self._overlay.show()
        self._overlay.raise_()
        self._overlay.activateWindow()

    def _finish_region_capture(self, logical_rect: QRect, pil_full):
        """
        Crop the pre-grabbed PIL image to the selected region.
        PIL's crop() works in the image's own pixel coordinates, which match
        the overlay's logical-pixel coordinate system exactly.
        """
        self._overlay = None
        cropped = pil_full.crop((
            logical_rect.x(),
            logical_rect.y(),
            logical_rect.x() + logical_rect.width(),
            logical_rect.y() + logical_rect.height(),
        ))
        self._save_card(cropped)

    # ═══════════════════════════════════════════════════════════════════════════
    # CAPTURE — MINIMIZE + FULL SCREEN
    # ═══════════════════════════════════════════════════════════════════════════

    def _minimize_capture(self):
        self._set_enabled(False)
        self.window().showMinimized()
        QTimer.singleShot(1500, self._after_minimize)

    def _after_minimize(self):
        self._save_card(self._pil_grab(*self._screen_bbox()))
        win = self.window()
        win.showNormal()
        win.raise_()
        win.activateWindow()
        self._set_enabled(True)

    def _set_enabled(self, on: bool):
        self.capture_btn.setEnabled(on)
        self.region_btn.setEnabled(on)
        self.minimize_btn.setEnabled(on)

    # ═══════════════════════════════════════════════════════════════════════════
    # SAVE & ADD CARD
    # ═══════════════════════════════════════════════════════════════════════════

    def _save_card(self, pil_img):
        """
        Save a PIL Image to disk as lossless PNG, then add a card to the gallery.
        PIL images preserve the full native pixel count regardless of Windows
        display scaling.
        """
        self._capture_counter += 1
        idx   = self._capture_counter
        stamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = f"screenshot_{idx:03d}_{stamp}.png"
        fpath = str(self._folder / fname)

        pil_img.save(fpath)          # lossless PNG at native resolution

        # Load the saved file as QPixmap for the card thumbnail
        display_px = QPixmap(fpath)

        # Fill the first pending placeholder card if one exists
        target = next((c for c in self._cards if c.is_placeholder), None)
        if target:
            target.set_captured(display_px, fpath)
            self._refresh_state()
            # Scroll to show the NEXT pending card so user sees what's next
            next_pend = next((c for c in self._cards if c.is_placeholder), None)
            if next_pend:
                QTimer.singleShot(
                    60, lambda w=next_pend: self._scroll.ensureWidgetVisible(w))
            else:
                # All done — scroll to bottom
                self._scroll.verticalScrollBar().setValue(
                    self._scroll.verticalScrollBar().maximum())
        else:
            # No placeholder — add a brand-new card
            card = ScreenshotCard(display_px, fpath, len(self._cards) + 1)
            if self._caption_queue:
                card.caption_edit.setText(self._caption_queue.pop(0))
            self._wire_card(card)
            self._cards.append(card)
            self._gallery_layout.insertWidget(
                self._gallery_layout.count() - 1, card)
            self._refresh_state()
            self._scroll.verticalScrollBar().setValue(
                self._scroll.verticalScrollBar().maximum())

    # ═══════════════════════════════════════════════════════════════════════════
    # GALLERY MANAGEMENT
    # ═══════════════════════════════════════════════════════════════════════════

    def _wire_card(self, card: ScreenshotCard):
        """Connect every card signal to its handler (call once per new card)."""
        card.sig_delete.connect(self._delete_card)
        card.sig_move_up.connect(self._move_card_up)
        card.sig_move_down.connect(self._move_card_down)
        card.sig_retake.connect(self._retake_card)
        card.sig_insert_below.connect(self._insert_card_below)

    def _retake_card(self, card: ScreenshotCard):
        """
        Reset a captured card back to a placeholder so it can be re-shot.
        The PNG that was saved to disk is preserved — only unlinked from the slot.
        Automatically switches to the Capture tab so the user can take the shot.
        """
        if card not in self._cards:
            return
        card.reset_to_placeholder()
        self._refresh_state()
        # Jump to Capture tab so the user can immediately shoot the replacement
        self._tabs.setCurrentIndex(1)
        # Scroll gallery so this pending card is in view when they switch back
        QTimer.singleShot(80, lambda: self._scroll.ensureWidgetVisible(card))

    def _insert_card_below(self, card: ScreenshotCard):
        """
        Insert a new blank placeholder card immediately after *card*.
        The user can type a caption for it and then capture the screenshot.
        """
        if card not in self._cards:
            return
        idx = self._cards.index(card)
        new_card = ScreenshotCard(None, "", idx + 2)   # 1-based visual index
        self._wire_card(new_card)
        # Insert in the internal list right after the reference card
        self._cards.insert(idx + 1, new_card)
        # Layout: position 0 = empty_lbl, so card[idx] is at layout pos idx+1,
        # inserting after it means layout pos = idx + 2
        self._gallery_layout.insertWidget(idx + 2, new_card)
        self._reindex()
        self._refresh_state()
        QTimer.singleShot(80, lambda: self._scroll.ensureWidgetVisible(new_card))

    def _delete_card(self, card: ScreenshotCard):
        if card not in self._cards:
            return
        self._cards.remove(card)
        self._gallery_layout.removeWidget(card)
        card.setParent(None)
        card.deleteLater()
        self._reindex()
        self._refresh_state()

    def _move_card_up(self, card: ScreenshotCard):
        i = self._cards.index(card)
        if i == 0:
            return
        self._cards[i], self._cards[i - 1] = self._cards[i - 1], self._cards[i]
        self._rebuild_layout()

    def _move_card_down(self, card: ScreenshotCard):
        i = self._cards.index(card)
        if i >= len(self._cards) - 1:
            return
        self._cards[i], self._cards[i + 1] = self._cards[i + 1], self._cards[i]
        self._rebuild_layout()

    def _rebuild_layout(self):
        for c in self._cards:
            self._gallery_layout.removeWidget(c)
        for i, c in enumerate(self._cards):
            c.set_index(i + 1)
            # Position: after empty_lbl (idx 0), before stretch (last)
            self._gallery_layout.insertWidget(i + 1, c)
        self._refresh_state()

    def _reindex(self):
        for i, c in enumerate(self._cards):
            c.set_index(i + 1)

    def _refresh_state(self):
        n_total    = len(self._cards)
        n_pending  = sum(1 for c in self._cards if c.is_placeholder)
        n_captured = n_total - n_pending

        self._empty_lbl.setVisible(n_total == 0)

        # Tab 3 — gallery header count
        if n_pending > 0:
            self._count_lbl.setText(
                f"{n_captured} captured  /  {n_total} total  "
                f"({n_pending} pending)")
        else:
            self._count_lbl.setText(
                f"{n_total} screenshot{'s' if n_total != 1 else ''}")

        # Tab 2 — prominent capture banner
        self._update_capture_banner(n_captured, n_total, n_pending)

    def _clear_all(self):
        if not self._cards:
            return
        if QMessageBox.question(
                self, "Clear All",
                "Remove all screenshots from the list?\n"
                "(Saved PNG files on disk are NOT deleted.)",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
        ) != QMessageBox.StandardButton.Yes:
            return
        self._force_clear()

    def _force_clear(self):
        for c in self._cards:
            self._gallery_layout.removeWidget(c)
            c.setParent(None)
            c.deleteLater()
        self._cards.clear()
        self._caption_queue   = []
        self._capture_counter = 0
        self._refresh_state()

    # ═══════════════════════════════════════════════════════════════════════════
    # SESSION — SAVE / LOAD
    # ═══════════════════════════════════════════════════════════════════════════

    def _save_session(self):
        """
        Serialise the current session to a JSON file (.srp.json).

        Stored data
        -----------
        version       : int   — file-format version for future compatibility
        report        : dict  — all report-settings fields
        cards         : list  — one entry per ScreenshotCard with:
                                  caption       str
                                  is_placeholder bool
                                  filepath      str  (absolute path to PNG or "")
        capture_counter : int — preserved so filenames don't collide on resume
        """
        # Default filename based on report title (or "session")
        title_slug = (
            self.title_edit.text().strip().replace(" ", "_") or "session"
        )
        default_path = str(self._folder / f"{title_slug}.srp.json")

        path, _ = QFileDialog.getSaveFileName(
            self, "Save Session",
            default_path,
            "Screen Report Session (*.srp.json);;All Files (*)"
        )
        if not path:
            return

        data = {
            "version": 1,
            "report": {
                "title":       self.title_edit.text(),
                "date":        self.date_edit.text(),
                "project":     self.project_edit.text(),
                "client":      self.client_edit.text(),
                "prepared_by": self.prepby_edit.text(),
                "location":    self.location_edit.text(),
                "save_folder": str(self._folder),
            },
            "capture_counter": self._capture_counter,
            "cards": [
                {
                    "caption":        c.get_caption(),
                    "is_placeholder": c.is_placeholder,
                    "filepath":       c.filepath or "",
                }
                for c in self._cards
            ],
        }

        try:
            Path(path).write_text(
                json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8"
            )
        except Exception as exc:
            QMessageBox.critical(
                self, "Save Failed",
                f"Could not write session file:\n{exc}"
            )
            return

        self._session_path = Path(path)
        QMessageBox.information(
            self, "Session Saved",
            f"Session saved to:\n{path}\n\n"
            f"Open this file with  Load Session  to resume later."
        )

    def _load_session(self):
        """
        Restore a previously saved .srp.json session.

        Cards whose PNG file is no longer on disk are restored as placeholders
        so the user can re-capture them — no data is silently lost.
        """
        start_dir = (
            str(self._session_path.parent)
            if self._session_path
            else str(self._folder)
        )
        path, _ = QFileDialog.getOpenFileName(
            self, "Load Session",
            start_dir,
            "Screen Report Session (*.srp.json);;All Files (*)"
        )
        if not path:
            return

        # ── Parse ─────────────────────────────────────────────────────────────
        try:
            raw = json.loads(Path(path).read_text(encoding="utf-8"))
        except Exception as exc:
            QMessageBox.critical(
                self, "Load Failed",
                f"Could not read session file:\n{exc}"
            )
            return

        if raw.get("version", 0) != 1:
            QMessageBox.warning(
                self, "Unknown Format",
                "This session file was created by a newer version of the app\n"
                "and may not load correctly."
            )

        # ── Confirm replacing existing session ────────────────────────────────
        if self._cards:
            ans = QMessageBox.question(
                self, "Replace Session?",
                "Loading a session will replace the current gallery.\n"
                "Saved PNG files on disk will NOT be deleted.\n\n"
                "Continue?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            )
            if ans != QMessageBox.StandardButton.Yes:
                return

        # ── Restore report settings ───────────────────────────────────────────
        rpt = raw.get("report", {})
        self.title_edit.setText(rpt.get("title", ""))
        self.date_edit.setText(rpt.get("date", ""))
        self.project_edit.setText(rpt.get("project", ""))
        self.client_edit.setText(rpt.get("client", ""))
        self.prepby_edit.setText(rpt.get("prepared_by", ""))
        self.location_edit.setText(rpt.get("location", ""))

        folder_str = rpt.get("save_folder", "")
        if folder_str and Path(folder_str).is_dir():
            self._folder = Path(folder_str)
            self.folder_edit.setText(str(self._folder))

        # ── Clear gallery and rebuild from cards list ─────────────────────────
        self._force_clear()          # wipes _cards and resets _capture_counter

        # Restore counter AFTER _force_clear() so it isn't zeroed out
        self._capture_counter = int(raw.get("capture_counter", 0))

        missing: list[str] = []
        for i, entry in enumerate(raw.get("cards", [])):
            caption  = entry.get("caption", "")
            was_cap  = not entry.get("is_placeholder", True)
            filepath = entry.get("filepath", "")
            file_ok  = was_cap and bool(filepath) and Path(filepath).is_file()

            if file_ok:
                card = ScreenshotCard(QPixmap(filepath), filepath, i + 1,
                                      caption=caption)
            else:
                # Pass None → constructor sets _is_placeholder = True automatically
                card = ScreenshotCard(None, "", i + 1, caption=caption)
                if was_cap and not file_ok:
                    missing.append(filepath or "(unknown)")

            self._wire_card(card)
            self._cards.append(card)
            self._gallery_layout.insertWidget(
                self._gallery_layout.count() - 1, card)

        self._session_path = Path(path)
        self._refresh_state()

        # ── Switch to Gallery tab so user can see what loaded ─────────────────
        self._tabs.setCurrentIndex(2)

        if missing:
            QMessageBox.warning(
                self, "Missing Files",
                f"{len(missing)} screenshot(s) could not be found on disk\n"
                "and have been restored as placeholders:\n\n"
                + "\n".join(missing[:10])
                + ("\n…" if len(missing) > 10 else "")
            )

    # ── Live preview ──────────────────────────────────────────────────────────

    def _toggle_live(self, on: bool):
        if on:
            self.live_btn.setText("  Live Preview  ON")
            self.live_btn.setIcon(qta.icon("fa5s.eye", color="#30b950"))
            self._live_timer.start()
        else:
            self.live_btn.setText("  Live Preview (1 fps)")
            self.live_btn.setIcon(qta.icon("fa5s.eye", color="#8b949e"))
            self._live_timer.stop()

    # ── Folder browse ──────────────────────────────────────────────────────────

    def _browse_folder(self):
        d = QFileDialog.getExistingDirectory(
            self, "Select Save Folder", str(self._folder))
        if d:
            self._folder = Path(d)
            self.folder_edit.setText(str(self._folder))

    # ═══════════════════════════════════════════════════════════════════════════
    # EXPORT — PDF
    # ═══════════════════════════════════════════════════════════════════════════

    def _export_pdf(self):
        if not any(not c.is_placeholder for c in self._cards):
            QMessageBox.warning(self, "No Screenshots",
                                "Capture at least one screenshot first.")
            return
        title = self.title_edit.text().strip() or "Screen Report"
        path, _ = QFileDialog.getSaveFileName(
            self, "Save PDF Report",
            str(self._folder / (title.replace(" ", "_") + ".pdf")),
            "PDF Files (*.pdf)")
        if not path:
            return
        try:
            self._build_pdf(path, title)
            QMessageBox.information(self, "Exported", f"PDF saved to:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Failed", str(exc))

    def _build_pdf(self, path: str, title: str):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.units import mm
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.colors import HexColor, white
        from reportlab.platypus import (
            SimpleDocTemplate, Paragraph, Spacer,
            HRFlowable, Image as RLImage, PageBreak,
            Table, TableStyle,
        )

        W, H = A4
        M    = 20 * mm

        c_green = HexColor("#1a6b2e")
        c_text  = HexColor("#1c1c1c")
        c_sub   = HexColor("#5a6370")
        c_hr    = HexColor("#d0d7de")
        c_row_a = HexColor("#f6f8fa")
        c_row_b = HexColor("#ffffff")

        def PS(name, **kw):
            return ParagraphStyle(name, **kw)

        sty_co   = PS("co",  fontSize=11, fontName="Helvetica",
                      textColor=c_sub, spaceAfter=2)
        sty_ti   = PS("ti",  fontSize=26, fontName="Helvetica-Bold",
                      textColor=c_text, spaceAfter=6)
        sty_h2   = PS("h2",  fontSize=13, fontName="Helvetica-Bold",
                      textColor=c_text, spaceAfter=4)
        sty_body = PS("bo",  fontSize=10, fontName="Helvetica",
                      textColor=c_text)
        sty_key  = PS("ke",  fontSize=10, fontName="Helvetica-Bold",
                      textColor=c_sub)
        sty_cap  = PS("ca",  fontSize=9,  fontName="Helvetica-Oblique",
                      textColor=c_sub, spaceAfter=4)
        sty_sl   = PS("sl",  fontSize=9,  fontName="Helvetica",
                      textColor=c_sub)
        sty_sln  = PS("sln", fontSize=9,  fontName="Helvetica",
                      textColor=c_text)

        doc = SimpleDocTemplate(
            path, pagesize=A4,
            leftMargin=M, rightMargin=M, topMargin=M, bottomMargin=M,
            title=title, author="Solareff Engineering")
        story = []

        # ── Cover ──────────────────────────────────────────────────────────────
        story.append(Paragraph("SOLAREFF ENGINEERING", sty_co))
        story.append(HRFlowable(width="100%", thickness=2.5,
                                color=c_green, spaceAfter=10))
        story.append(Paragraph(title, sty_ti))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=c_hr, spaceAfter=14))

        def kv(k, v):
            return [Paragraph(k, sty_key), Paragraph(v or "—", sty_body)]

        meta = Table(
            [kv("Date",          self.date_edit.text().strip()),
             kv("Project No.",   self.project_edit.text().strip()),
             kv("Client / Site", self.client_edit.text().strip()),
             kv("Prepared By",   self.prepby_edit.text().strip()),
             kv("Location",      self.location_edit.text().strip())],
            colWidths=[42 * mm, W - 2 * M - 42 * mm])
        meta.setStyle(TableStyle([
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [c_row_a, c_row_b]),
            ("LEFTPADDING",    (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",   (0, 0), (-1, -1), 6),
            ("TOPPADDING",     (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 5),
            ("GRID",           (0, 0), (-1, -1), 0.5, c_hr),
            ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(meta)
        story.append(Spacer(1, 18))

        # ── Table of contents ─────────────────────────────────────────────────
        story.append(Paragraph("Contents", sty_h2))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=c_hr, spaceAfter=6))
        captured = [c for c in self._cards if not c.is_placeholder]

        toc_data = [
            [Paragraph("#", sty_key),
             Paragraph("Caption / Sub-heading", sty_key)],
        ]
        for i, c in enumerate(captured, 1):
            toc_data.append([Paragraph(str(i), sty_body),
                             Paragraph(c.get_caption(), sty_body)])
        toc = Table(toc_data,
                    colWidths=[14 * mm, W - 2 * M - 14 * mm])
        toc.setStyle(TableStyle([
            ("BACKGROUND",     (0, 0), (-1, 0), c_green),
            ("TEXTCOLOR",      (0, 0), (-1, 0), white),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [c_row_a, c_row_b]),
            ("LEFTPADDING",    (0, 0), (-1, -1), 6),
            ("RIGHTPADDING",   (0, 0), (-1, -1), 6),
            ("TOPPADDING",     (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING",  (0, 0), (-1, -1), 4),
            ("GRID",           (0, 0), (-1, -1), 0.5, c_hr),
            ("VALIGN",         (0, 0), (-1, -1), "MIDDLE"),
        ]))
        story.append(toc)
        story.append(Spacer(1, 20))

        # ── Sign-off ───────────────────────────────────────────────────────────
        story.append(Paragraph("Sign-off", sty_h2))
        story.append(HRFlowable(width="100%", thickness=0.5,
                                color=c_hr, spaceAfter=10))
        line_long  = "_" * 44
        line_short = "_" * 22
        sig_rows = [
            ["Prepared by:",  line_long, "Date:", line_short],
            ["Reviewed by:",  line_long, "Date:", line_short],
            ["Approved by:",  line_long, "Date:", line_short],
        ]
        sig_tbl = Table(
            [[Paragraph(c, sty_sl if i % 2 == 0 else sty_sln)
              for i, c in enumerate(row)]
             for row in sig_rows],
            colWidths=[28 * mm, 65 * mm, 16 * mm, 58 * mm])
        sig_tbl.setStyle(TableStyle([
            ("TOPPADDING",    (0, 0), (-1, -1), 14),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("VALIGN",        (0, 0), (-1, -1), "BOTTOM"),
        ]))
        story.append(sig_tbl)
        story.append(PageBreak())

        # ── Image pages ────────────────────────────────────────────────────────
        usable_w = W - 2 * M
        usable_h = H - 2 * M - 50 * mm

        for i, card in enumerate(captured, 1):
            story.append(Paragraph(f"{i}.  {card.get_caption()}", sty_h2))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=c_hr, spaceAfter=8))
            try:
                img = RLImage(card.filepath)
                iw, ih = img.imageWidth, img.imageHeight
                scale  = min(usable_w / iw, usable_h / ih, 1.0)
                img.drawWidth  = iw * scale
                img.drawHeight = ih * scale
                story.append(img)
            except Exception:
                story.append(Paragraph(
                    f"[Could not load: {card.filepath}]", sty_cap))
            story.append(Spacer(1, 6))
            story.append(Paragraph(
                f"Figure {i}:  {card.get_caption()}", sty_cap))
            if i < len(captured):
                story.append(PageBreak())

        doc.build(story)

    # ═══════════════════════════════════════════════════════════════════════════
    # EXPORT — WORD
    # ═══════════════════════════════════════════════════════════════════════════

    def _export_word(self):
        if not any(not c.is_placeholder for c in self._cards):
            QMessageBox.warning(self, "No Screenshots",
                                "Capture at least one screenshot first.")
            return
        title = self.title_edit.text().strip() or "Screen Report"
        path, _ = QFileDialog.getSaveFileName(
            self, "Save Word Report",
            str(self._folder / (title.replace(" ", "_") + ".docx")),
            "Word Documents (*.docx)")
        if not path:
            return
        try:
            self._build_word(path, title)
            QMessageBox.information(self, "Exported",
                                    f"Word document saved to:\n{path}")
        except Exception as exc:
            QMessageBox.critical(self, "Export Failed", str(exc))

    def _build_word(self, path: str, title: str):
        from docx import Document
        from docx.shared import Inches, Pt, Cm, RGBColor
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        GREEN = RGBColor(0x1a, 0x6b, 0x2e)
        GREY  = RGBColor(0x5a, 0x63, 0x70)

        doc = Document()
        for sec in doc.sections:
            sec.top_margin    = Cm(2.0)
            sec.bottom_margin = Cm(2.0)
            sec.left_margin   = Cm(2.5)
            sec.right_margin  = Cm(2.5)

        def add_hr(colour="238636"):
            p_ = doc.add_paragraph()
            pPr = p_._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot  = OxmlElement("w:bottom")
            bot.set(qn("w:val"),   "single")
            bot.set(qn("w:sz"),    "8")
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), colour)
            pBdr.append(bot)
            pPr.append(pBdr)
            return p_

        # Header
        p = doc.add_paragraph("SOLAREFF ENGINEERING")
        p.runs[0].font.size = Pt(11)
        p.runs[0].font.color.rgb = GREY
        add_hr()

        h = doc.add_heading(title, level=1)
        for run in h.runs:
            run.font.size = Pt(22)
        add_hr("d0d7de")

        # Meta table
        mt = doc.add_table(rows=0, cols=2)
        mt.style = "Table Grid"
        for key, val in [
            ("Date",          self.date_edit.text().strip()),
            ("Project No.",   self.project_edit.text().strip()),
            ("Client / Site", self.client_edit.text().strip()),
            ("Prepared By",   self.prepby_edit.text().strip()),
            ("Location",      self.location_edit.text().strip()),
        ]:
            row = mt.add_row()
            row.cells[0].width = Cm(4.5)
            row.cells[0].text  = key
            row.cells[0].paragraphs[0].runs[0].font.bold      = True
            row.cells[0].paragraphs[0].runs[0].font.color.rgb = GREY
            row.cells[1].text = val or "—"

        doc.add_paragraph()

        captured = [c for c in self._cards if not c.is_placeholder]

        # TOC
        doc.add_heading("Contents", level=2)
        toc = doc.add_table(rows=1, cols=2)
        toc.style = "Table Grid"
        for c, t in zip(toc.rows[0].cells, ["#", "Caption / Sub-heading"]):
            c.text = t
            c.paragraphs[0].runs[0].font.bold = True
        for i, card in enumerate(captured, 1):
            row = toc.add_row().cells
            row[0].text = str(i)
            row[1].text = card.get_caption()
        doc.add_paragraph()

        # Sign-off
        doc.add_heading("Sign-off", level=2)
        sig = doc.add_table(rows=3, cols=4)
        sig.style = "Table Grid"
        for r_idx, lbl in enumerate(
                ["Prepared by:", "Reviewed by:", "Approved by:"]):
            cells = sig.rows[r_idx].cells
            cells[0].text = lbl
            cells[0].paragraphs[0].runs[0].font.color.rgb = GREY
            cells[1].text = " " * 40
            cells[2].text = "Date:"
            cells[2].paragraphs[0].runs[0].font.color.rgb = GREY
            cells[3].text = " " * 20
            sig.rows[r_idx].height = Cm(1.6)

        doc.add_page_break()

        # Image pages — placeholders are skipped
        for i, card in enumerate(captured, 1):
            doc.add_heading(f"{i}.  {card.get_caption()}", level=2)
            add_hr("d0d7de")
            try:
                doc.add_picture(card.filepath, width=Inches(6.0))
                cap = doc.add_paragraph(
                    f"Figure {i}:  {card.get_caption()}")
                cap.runs[0].font.italic    = True
                cap.runs[0].font.size      = Pt(9)
                cap.runs[0].font.color.rgb = GREY
            except Exception:
                doc.add_paragraph(f"[Could not load: {card.filepath}]")
            if i < len(captured):
                doc.add_page_break()

        doc.save(path)

    # ═══════════════════════════════════════════════════════════════════════════
    # Template / caption queue
    # ═══════════════════════════════════════════════════════════════════════════

    def _show_template_dialog(self):
        dlg = _TemplateDialog(
            current_queue_len=len(self._caption_queue), parent=self)
        if dlg.exec() != QDialog.DialogCode.Accepted:
            return
        caps = dlg.get_captions()
        if not caps:
            return
        self._apply_template_as_placeholders(caps)

    def _apply_template_as_placeholders(self, captions: list):
        """
        Populate the gallery with placeholder cards for every caption in
        *captions*.  If the gallery already has content the user is asked
        whether to clear it first.
        """
        if self._cards:
            n_cap  = sum(1 for c in self._cards if not c.is_placeholder)
            n_pend = sum(1 for c in self._cards if c.is_placeholder)
            msg = (f"Load {len(captions)} caption placeholders into the gallery?\n\n"
                   f"Current gallery:  {n_cap} captured,  {n_pend} pending.\n\n"
                   "Clear the gallery and start fresh?\n"
                   "(Saved PNG files on disk are not deleted.)")
            if QMessageBox.question(
                self, "Load Template", msg,
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No,
            ) != QMessageBox.StandardButton.Yes:
                return
            self._force_clear()

        for caption in captions:
            card = ScreenshotCard(
                None, "", len(self._cards) + 1, caption=caption)
            self._wire_card(card)
            self._cards.append(card)
            self._gallery_layout.insertWidget(
                self._gallery_layout.count() - 1, card)

        self._caption_queue = []   # placeholders replace the text queue
        self._reindex()
        self._refresh_state()
        # Scroll to top so the first pending card is visible
        QTimer.singleShot(60, lambda: self._scroll.verticalScrollBar().setValue(0))

    def _update_capture_banner(self, n_captured: int = -1,
                               n_total: int = -1, n_pending: int = -1):
        """
        Update the prominent status banner on the Capture tab.
        Shows the next pending caption and overall progress.
        """
        if n_captured < 0:
            n_total    = len(self._cards)
            n_pending  = sum(1 for c in self._cards if c.is_placeholder)
            n_captured = n_total - n_pending

        next_pend = next((c for c in self._cards if c.is_placeholder), None)
        if next_pend:
            cap = next_pend.get_caption()
            if len(cap) > 60:
                cap = cap[:58] + "…"
            self._next_lbl.setText(f"▶   {cap}")
            self._next_lbl.setStyleSheet(
                "color: #58a6ff; font-size: 14pt; font-weight: bold;")
            self._progress_lbl.setText(
                f"{n_captured} captured  ·  {n_total} total  ·  {n_pending} pending")
        elif n_total > 0:
            self._next_lbl.setText("✓   All screenshots captured")
            self._next_lbl.setStyleSheet(
                "color: #3fb950; font-size: 14pt; font-weight: bold;")
            self._progress_lbl.setText(
                f"{n_captured} screenshot{'s' if n_captured != 1 else ''}"
                f"  —  switch to  Gallery & Export  to generate the report")
        else:
            self._next_lbl.setText("Ready to capture")
            self._next_lbl.setStyleSheet(
                "color: #8b949e; font-size: 13pt;")
            self._progress_lbl.setText(
                "Load a template (Gallery & Export tab) or capture freely")

    # ═══════════════════════════════════════════════════════════════════════════
    # BaseCalculator interface
    # ═══════════════════════════════════════════════════════════════════════════

    def get_inputs(self) -> dict:
        return {
            "Report Title":     self.title_edit.text(),
            "Date":             self.date_edit.text(),
            "Project No.":      self.project_edit.text(),
            "Client / Site":    self.client_edit.text(),
            "Prepared By":      self.prepby_edit.text(),
            "Location":         self.location_edit.text(),
            "Save Folder":      str(self._folder),
        }

    def get_results(self) -> dict:
        return {"Screenshots Captured": str(len(self._cards))}

    def reset(self):
        self.title_edit.clear()
        self.project_edit.clear()
        self.client_edit.clear()
        self.prepby_edit.clear()
        self.location_edit.clear()
        self.date_edit.setText(
            datetime.date.today().strftime("%d %B %Y"))
        self._force_clear()
